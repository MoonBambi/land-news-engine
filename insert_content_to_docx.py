"""
Step 1: 用序言.md的章/节标题替换输入docx中的旧标题
Step 2: 在有正文内容的小标题下插入正文段落
Step 3: 将【待插入：...】占位符替换为 imgs/ 目录下的实际图片

Usage:
  python insert_content_to_docx.py                              # 默认 初稿0.1.docx -> 初稿0.2.docx
  python insert_content_to_docx.py --input 初稿0.2_new.docx --output 初稿0.3.docx
"""
import argparse
import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

BASE_DIR = Path(__file__).parent
DEFAULT_SRC = BASE_DIR / "初稿0.1.docx"
MD_PATH = BASE_DIR / "序言.md"
DEFAULT_OUT = BASE_DIR / "初稿0.2.docx"
FALLBACK = BASE_DIR / "初稿0.2_new.docx"
IMGS_DIR = BASE_DIR / "imgs"

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def q(name):
    return qn(f"w:{name}")


# ── Markdown Parser ──

def parse_md(md_path):
    """
    Returns {
        "preface_paragraphs": [str, ...],   # 序言正文
        "chapters": [{num, title, sections: [{number, title, level, paragraphs}]}]
    }
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    skip_titles = {"论文目录：", "摘 要", "Abstract", "参考文献", "致 谢", "附 录"}

    chapters = []
    current_chapter = None
    current_section = None
    preface_paras = []

    # 序言 content lives between the "序言" line and the first content heading
    in_preface_body = False
    preface_lines_buffer = []

    for line in lines:
        stripped = line.rstrip('\n').strip()

        # Detect "序言" heading
        if stripped == "序言" and not chapters:
            in_preface_body = True
            continue

        # If we were in preface, collect body lines until a heading appears
        if in_preface_body:
            ch_test = re.match(r'^(第\s*\d+\s*章(?![:：])|1\.\d+|论文目录|摘\s*要|Abstract)', stripped)
            if ch_test or any(stripped.startswith(t) for t in skip_titles):
                # End of preface body
                in_preface_body = False
                preface_paras = _group_paragraphs(preface_lines_buffer)

            if not stripped:
                preface_lines_buffer.append("")
            else:
                preface_lines_buffer.append(stripped)
            continue

        if not stripped:
            continue

        if any(stripped.startswith(t) for t in skip_titles):
            current_section = None
            continue

        # Chapter heading
        # Chapter heading (must NOT have colon after chapter title)
        ch_match = re.match(r'^第\s*(\d+)\s*章(?![:：])\s*(.*)', stripped)
        if ch_match:
            current_chapter = {
                'num': int(ch_match.group(1)),
                'title': stripped,
                'sections': []
            }
            chapters.append(current_chapter)
            current_section = None
            continue

        # Section heading
        sec_match = re.match(r'^(\d+\.\d+(?:\.\d+)?)\s+(.*)', stripped)
        if sec_match:
            num = sec_match.group(1)
            title = sec_match.group(2).strip()
            current_section = {
                'number': num,
                'title': title,
                'level': num.count('.'),
                'paragraphs': []
            }
            if current_chapter:
                current_chapter['sections'].append(current_section)
            else:
                # Standalone section (e.g. 1.1 before chapter heading)
                chapters.append({
                    'num': 0,
                    'title': f'{num} {title}',
                    'sections': [current_section]
                })
            continue

        # Body text
        if current_section is not None:
            current_section['paragraphs'].append(stripped)

    # Flush remaining preface
    if in_preface_body and preface_lines_buffer:
        preface_paras = _group_paragraphs(preface_lines_buffer)

    return {
        'preface_paragraphs': preface_paras,
        'chapters': chapters,
    }


def _group_paragraphs(lines):
    """Merge consecutive non-empty lines into paragraphs."""
    result = []
    buf = []
    for line in lines:
        if not line.strip():
            if buf:
                result.append(" ".join(buf))
                buf = []
        else:
            buf.append(line.strip())
    if buf:
        result.append(" ".join(buf))
    return result


# ── docx Element Builders ──

def _strip_sect_pr(elem):
    pPr = elem.find(q("pPr"))
    if pPr is not None:
        sp = pPr.find(q("sectPr"))
        if sp is not None:
            pPr.remove(sp)


def _disable_keep_next(elem):
    pPr = elem.find(q("pPr"))
    if pPr is None:
        pPr = etree.SubElement(elem, q("pPr"))
        elem.insert(0, pPr)
    kn = pPr.find(q("keepNext"))
    if kn is None:
        kn = etree.SubElement(pPr, q("keepNext"))
    kn.set(q("val"), '0')


def _make_heading_para(text, style_val, template_elem=None):
    """Clone a heading paragraph from template or build from scratch."""
    if template_elem is not None:
        p = etree.fromstring(etree.tostring(template_elem))
        _strip_sect_pr(p)
    else:
        p = etree.Element(q("p"))
        pPr = etree.SubElement(p, q("pPr"))
        etree.SubElement(pPr, q("pStyle")).set(q("val"), style_val)

    # Set/override style
    pPr = p.find(q("pPr"))
    if pPr is None:
        pPr = etree.SubElement(p, q("pPr"))
    ps = pPr.find(q("pStyle"))
    if ps is None:
        ps = etree.SubElement(pPr, q("pStyle"))
    ps.set(q("val"), style_val)

    # Clear old runs, keep first run properties
    old_runs = list(p.findall(q("r")))
    first_rPr = None
    if old_runs:
        fr = old_runs[0].find(q("rPr"))
        if fr is not None:
            first_rPr = etree.fromstring(etree.tostring(fr))
    for r_elem in old_runs:
        p.remove(r_elem)

    # Create text run
    r_elem = etree.SubElement(p, q("r"))
    if first_rPr is not None:
        r_elem.append(first_rPr)
    else:
        etree.SubElement(r_elem, q("rPr"))
    t_elem = etree.SubElement(r_elem, q("t"))
    t_elem.text = text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    _disable_keep_next(p)
    return p


def _make_body_para(text):
    """Create a body-text paragraph (Normal style, 宋体 小四, 1.5x line spacing, indent)."""
    p = etree.Element(q("p"))
    pPr = etree.SubElement(p, q("pPr"))
    pStyle = etree.SubElement(pPr, q("pStyle"))
    pStyle.set(q("val"), "Normal")

    spacing = etree.SubElement(pPr, q("spacing"))
    spacing.set(q("line"), "360")
    spacing.set(q("lineRule"), "auto")

    ind = etree.SubElement(pPr, q("ind"))
    ind.set(q("firstLine"), "480")

    r = etree.SubElement(p, q("r"))
    rPr = etree.SubElement(r, q("rPr"))
    rFonts = etree.SubElement(rPr, q("rFonts"))
    rFonts.set(q("eastAsia"), "宋体")
    rFonts.set(q("ascii"), "Times New Roman")
    rFonts.set(q("hAnsi"), "Times New Roman")
    sz = etree.SubElement(rPr, q("sz"))
    sz.set(q("val"), "24")
    szCs = etree.SubElement(rPr, q("szCs"))
    szCs.set(q("val"), "24")

    t = etree.SubElement(r, q("t"))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    return p


def _is_table_row(text):
    """Detect a markdown table row (starts and ends with |)."""
    return text.strip().startswith("|") and text.strip().endswith("|")


def _is_table_separator(text):
    """Detect a markdown table separator row like |---|---|."""
    stripped = text.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    inner = stripped[1:-1]
    return all(cell.strip() == "" or all(ch in "-:" for ch in cell.strip()) for cell in inner.split("|"))


def _parse_table_cells(line):
    """Extract cell texts from a markdown table row."""
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _make_table_cell_para(text, font_east_asia="楷体", font_ascii="Times New Roman", sz_val="21"):
    """Create a table cell paragraph (5号 = 10.5pt = 21 half-points, 楷体, centered)."""
    p = etree.Element(q("p"))
    pPr = etree.SubElement(p, q("pPr"))
    # Center justify
    jc = etree.SubElement(pPr, q("jc"))
    jc.set(q("val"), "center")
    # No indent for table cells
    ind = etree.SubElement(pPr, q("ind"))
    ind.set(q("firstLine"), "0")

    r = etree.SubElement(p, q("r"))
    rPr = etree.SubElement(r, q("rPr"))
    rFonts = etree.SubElement(rPr, q("rFonts"))
    rFonts.set(q("eastAsia"), font_east_asia)
    rFonts.set(q("ascii"), font_ascii)
    rFonts.set(q("hAnsi"), font_ascii)
    sz = etree.SubElement(rPr, q("sz"))
    sz.set(q("val"), sz_val)  # 5号 = 10.5pt = 21 half-points
    szCs = etree.SubElement(rPr, q("szCs"))
    szCs.set(q("val"), sz_val)

    t = etree.SubElement(r, q("t"))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p


def _make_table_caption_para(text):
    """Create a centered table caption paragraph (5号, 楷体)."""
    p = etree.Element(q("p"))
    pPr = etree.SubElement(p, q("pPr"))
    pStyle = etree.SubElement(pPr, q("pStyle"))
    pStyle.set(q("val"), "Normal")
    jc = etree.SubElement(pPr, q("jc"))
    jc.set(q("val"), "center")
    ind = etree.SubElement(pPr, q("ind"))
    ind.set(q("firstLine"), "0")

    r = etree.SubElement(p, q("r"))
    rPr = etree.SubElement(r, q("rPr"))
    rFonts = etree.SubElement(rPr, q("rFonts"))
    rFonts.set(q("eastAsia"), "楷体")
    rFonts.set(q("ascii"), "Times New Roman")
    rFonts.set(q("hAnsi"), "Times New Roman")
    sz = etree.SubElement(rPr, q("sz"))
    sz.set(q("val"), "21")
    szCs = etree.SubElement(rPr, q("szCs"))
    szCs.set(q("val"), "21")

    t = etree.SubElement(r, q("t"))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p


def _make_docx_table(data_rows, num_cols):
    """
    Create a w:tbl element from markdown table data_rows (list of cell lists).
    Table cells use 5号楷体 centered formatting.
    """
    tbl = etree.Element(q("tbl"))
    tblPr = etree.SubElement(tbl, q("tblPr"))
    tblStyle = etree.SubElement(tblPr, q("tblStyle"))
    tblStyle.set(q("val"), "TableGrid")
    tblW = etree.SubElement(tblPr, q("tblW"))
    tblW.set(q("w"), "5000")
    tblW.set(q("type"), "pct")
    jc = etree.SubElement(tblPr, q("jc"))
    jc.set(q("val"), "center")

    tblGrid = etree.SubElement(tbl, q("tblGrid"))
    for _ in range(num_cols):
        gridCol = etree.SubElement(tblGrid, q("gridCol"))
        gridCol.set(q("w"), str(9072 // max(num_cols, 1)))

    for row_idx, cells in enumerate(data_rows):
        tr = etree.SubElement(tbl, q("tr"))
        for cell_text in cells:
            tc = etree.SubElement(tr, q("tc"))
            tcPr = etree.SubElement(tc, q("tcPr"))
            tcW = etree.SubElement(tcPr, q("tcW"))
            tcW.set(q("w"), str(9072 // max(num_cols, 1)))
            tcW.set(q("type"), "dxa")
            tcBorders = etree.SubElement(tcPr, q("tcBorders"))
            for edge in ("top", "left", "bottom", "right"):
                border = etree.SubElement(tcBorders, q(edge))
                border.set(q("val"), "single")
                border.set(q("sz"), "4")
                border.set(q("space"), "0")
                border.set(q("color"), "000000")
            # Header row background
            if row_idx == 0:
                tcShd = etree.SubElement(tcPr, q("shd"))
                tcShd.set(q("val"), "clear")
                tcShd.set(q("color"), "auto")
                tcShd.set(q("fill"), "D9E2F3")

            para = _make_table_cell_para(cell_text,
                                         font_east_asia="楷体" if row_idx > 0 else "黑体",
                                         sz_val="21" if row_idx > 0 else "21")
            tc.append(para)

    return tbl


def _process_paragraphs_to_elements(paragraphs):
    """
    Convert a list of paragraph strings to docx elements.
    Detects markdown tables and converts them to proper w:tbl elements.
    Returns list of (is_table, element) tuples.
    """
    elements = []
    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]

        # If this line is a table row, collect consecutive table rows
        if _is_table_row(text) and not _is_table_separator(text):
            # Check if previous line was a table caption (not a | row, not a heading)
            caption_text = None
            caption_idx = i - 1
            if caption_idx >= 0:
                prev = paragraphs[caption_idx]
                if not _is_table_row(prev) and not prev.strip().startswith("#"):
                    # Use it as caption
                    caption_text = prev
                    elements.pop()  # Remove caption paragraph if already added
                    if elements and elements[-1][0] is False:
                        pass  # Will replace with caption

            # Collect all data rows (non-separator | lines)
            data_rows = []
            while i < len(paragraphs):
                line = paragraphs[i]
                if _is_table_row(line):
                    if not _is_table_separator(line):
                        cells = _parse_table_cells(line)
                        data_rows.append(cells)
                    i += 1
                else:
                    break

            if data_rows:
                num_cols = max(len(row) for row in data_rows)
                # Pad rows to same column count
                for row in data_rows:
                    while len(row) < num_cols:
                        row.append("")
                # Trim extra columns
                data_rows = [row[:num_cols] for row in data_rows]

                # Add caption as centered paragraph before table
                if caption_text:
                    elements.append((False, _make_table_caption_para(caption_text)))

                tbl = _make_docx_table(data_rows, num_cols)
                elements.append((True, tbl))
        else:
            elements.append((False, _make_body_para(text)))
            i += 1

    return elements


# ── Image Insertion ──

def _scan_images():
    """Scan imgs/ dir and build {figure_number: path} mapping."""
    mapping = {}
    if not IMGS_DIR.exists():
        return mapping
    for f in IMGS_DIR.iterdir():
        if not f.suffix.lower() == '.png':
            continue
        stem = f.stem
        # Try "图 2-1 ..." or "3-1" pattern
        m = re.match(r'图\s*(\d+[-.]\d+)', stem)
        if not m:
            m = re.match(r'(\d+[-.]\d+)', stem)
        if m:
            num = m.group(1).replace('.', '-')
            mapping[num] = str(f)
    return mapping


def _make_image_caption_para(text):
    """Create a centered image caption paragraph (5号楷体)."""
    p = etree.Element(q("p"))
    pPr = etree.SubElement(p, q("pPr"))
    jc = etree.SubElement(pPr, q("jc"))
    jc.set(q("val"), "center")
    ind = etree.SubElement(pPr, q("ind"))
    ind.set(q("firstLine"), "0")

    r = etree.SubElement(p, q("r"))
    rPr = etree.SubElement(r, q("rPr"))
    rFonts = etree.SubElement(rPr, q("rFonts"))
    rFonts.set(q("eastAsia"), "楷体")
    rFonts.set(q("ascii"), "Times New Roman")
    rFonts.set(q("hAnsi"), "Times New Roman")
    sz = etree.SubElement(rPr, q("sz"))
    sz.set(q("val"), "21")
    szCs = etree.SubElement(rPr, q("szCs"))
    szCs.set(q("val"), "21")

    t = etree.SubElement(r, q("t"))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p


def _insert_images(doc):
    """
    Find 【待插入：...】 placeholder paragraphs in the docx,
    match against imgs/ PNG files, and replace placeholders with real images.
    """
    image_map = _scan_images()
    if not image_map:
        print("\n  imgs/ 目录为空或无 PNG 文件，跳过图片插入")
        return

    print(f"\n=== 5. 插入图片 (imgs/) ===")
    print(f"  扫描到 {len(image_map)} 个 PNG 文件")

    body = doc.element.body
    # Collect all paragraph elements
    all_paras = list(doc.paragraphs)
    inserted = 0
    skipped = []

    for i, para in enumerate(all_paras):
        text = (para.text or '').strip()
        if not text.startswith('【待插入：'):
            continue

        # Extract figure number from the previous paragraph (caption line like "图 2-1 ...")
        fig_num = None
        caption_text = ""
        if i > 0:
            prev_para = all_paras[i - 1]
            prev_text = (prev_para.text or '').strip()
            m = re.match(r'图\s*(\d+[-.]\d+)\s+(.*)', prev_text)
            if m:
                fig_num = m.group(1).replace('.', '-')
                caption_text = m.group(2).strip()

        if not fig_num:
            skipped.append(text[:50])
            continue

        img_path = image_map.get(fig_num)
        if not img_path:
            # Try loose match
            for key, path in image_map.items():
                if fig_num in key or key in fig_num:
                    img_path = path
                    break
        if not img_path:
            skipped.append(fig_num)
            continue

        # Get the paragraph element of the previous paragraph (caption) and current (placeholder)
        try:
            prev_elem = prev_para._element
            curr_elem = para._element
        except AttributeError:
            skipped.append(f"{fig_num} (no element)")
            continue

        # Create image paragraph (centered)
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        # Insert image at 80% of page width (about 14cm for A4)
        run.add_picture(img_path, width=Cm(14))

        # Insert image paragraph after the placeholder, then remove caption + placeholder
        curr_elem.addnext(img_para._element)

        # Move caption text to after image
        caption_para_elem = _make_image_caption_para(f"图 {fig_num} {caption_text}")
        img_para._element.addnext(caption_para_elem)

        # Remove the original caption paragraph and placeholder
        body.remove(prev_elem)
        body.remove(curr_elem)

        # Clean up the empty paragraph added by doc.add_paragraph()
        empty_p = img_para._element.getprevious()
        if empty_p is not None and empty_p.tag == q("p"):
            texts = []
            for t_elem in empty_p.iter(q("t")):
                if t_elem.text:
                    texts.append(t_elem.text)
            if not ''.join(texts).strip():
                body.remove(empty_p)

        inserted += 1
        print(f"  [OK] 图 {fig_num} {caption_text} <- {Path(img_path).name}")

    if skipped:
        print(f"  [SKIP] 未匹配图片: {skipped}")
    print(f"  共插入 {inserted} 张图片")


def _format_references(doc):
    """
    Format the 参考文献 section per thesis standard (GB/T 7714):
    - Heading: new page, 3号黑体, centered
    - Entries: 小4号宋体, [N] numbered, top-aligned
    """
    body = doc.element.body
    refs_head = None

    # Find "参考文献" heading paragraph (could be style 1 or un-styled)
    for child in body.iter(q("p")):
        texts = []
        for t_elem in child.iter(q("t")):
            if t_elem.text:
                texts.append(t_elem.text)
        full = ''.join(texts).strip()
        if full == '参考文献':
            refs_head = child
            break

    if refs_head is None:
        print('\n  未找到参考文献标题，跳过参考文献格式化')
        return

    print("\n=== 6. 格式化参考文献 ===")

    # 1) Page break before 参考文献 heading
    pPr = refs_head.find(q("pPr"))
    if pPr is None:
        pPr = etree.SubElement(refs_head, q("pPr"))
        refs_head.insert(0, pPr)
    # Remove any existing pageBreakBefore, then add
    for pb in list(pPr.findall(q("pageBreakBefore"))):
        pPr.remove(pb)
    pb_elem = etree.SubElement(pPr, q("pageBreakBefore"))
    pb_elem.set(q("val"), "true")

    # 2) Heading format: 3号黑体 (16pt = 32 half-pts), centered
    # Remove existing style reference (it might be '1' or other)
    ps_elem = pPr.find(q("pStyle"))
    if ps_elem is not None:
        ps_elem.set(q("val"), "Normal")
    else:
        ps_elem = etree.SubElement(pPr, q("pStyle"))
        ps_elem.set(q("val"), "Normal")

    # Center
    jc = pPr.find(q("jc"))
    if jc is None:
        jc = etree.SubElement(pPr, q("jc"))
    jc.set(q("val"), "center")

    # Update all runs to 3号黑体
    for r_elem in refs_head.findall(q("r")):
        rPr = r_elem.find(q("rPr"))
        if rPr is None:
            rPr = etree.SubElement(r_elem, q("rPr"))
            r_elem.insert(0, rPr)
        rFonts = rPr.find(q("rFonts"))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, q("rFonts"))
        rFonts.set(q("eastAsia"), "黑体")
        rFonts.set(q("ascii"), "Times New Roman")
        rFonts.set(q("hAnsi"), "Times New Roman")
        sz = rPr.find(q("sz"))
        if sz is None:
            sz = etree.SubElement(rPr, q("sz"))
        sz.set(q("val"), "32")  # 3号 = 16pt = 32
        szCs = rPr.find(q("szCs"))
        if szCs is None:
            szCs = etree.SubElement(rPr, q("szCs"))
        szCs.set(q("val"), "32")

    count = 0
    # 3) Format [N] reference entries: 小4号宋体 (12pt), flush-left
    sibling = refs_head.getnext()
    while sibling is not None:
        # Stop at next heading-style paragraph or 致谢
        if sibling.tag == q("p"):
            texts = []
            for t_elem in sibling.iter(q("t")):
                if t_elem.text:
                    texts.append(t_elem.text)
            full = ''.join(texts).strip()
            if full == '致 谢' or full == '致谢':
                break
            # Check if this is a heading style
            pPr2 = sibling.find(q("pPr"))
            if pPr2 is not None:
                ps = pPr2.find(q("pStyle"))
                if ps is not None and ps.get(q("val")) in ('1', '2'):
                    break

            # Format as reference entry
            # Remove indent
            ind = pPr2.find(q("ind"))
            if ind is None:
                ind = etree.SubElement(pPr2, q("ind"))
            ind.set(q("firstLine"), "0")
            ind.set(q("left"), "0")

            # Update all runs to 小4号宋体
            for r_elem in sibling.findall(q("r")):
                rPr3 = r_elem.find(q("rPr"))
                if rPr3 is None:
                    rPr3 = etree.SubElement(r_elem, q("rPr"))
                    r_elem.insert(0, rPr3)
                rFonts = rPr3.find(q("rFonts"))
                if rFonts is None:
                    rFonts = etree.SubElement(rPr3, q("rFonts"))
                rFonts.set(q("eastAsia"), "宋体")
                rFonts.set(q("ascii"), "Times New Roman")
                rFonts.set(q("hAnsi"), "Times New Roman")
                sz = rPr3.find(q("sz"))
                if sz is None:
                    sz = etree.SubElement(rPr3, q("sz"))
                sz.set(q("val"), "24")  # 小4号 = 12pt = 24
                szCs = rPr3.find(q("szCs"))
                if szCs is None:
                    szCs = etree.SubElement(rPr3, q("szCs"))
                szCs.set(q("val"), "24")

            count += 1

        sibling = sibling.getnext()

    print(f"  格式化 {count} 条参考文献条目")
    print(f"  标题: 3号黑体居中 + 分页")


def _add_chapter_page_breaks(doc):
    """Add page breaks before each chapter heading (第X章)."""
    body = doc.element.body
    count = 0
    for child in body.iter(q("p")):
        pPr = child.find(q("pPr"))
        if pPr is None:
            continue
        ps = pPr.find(q("pStyle"))
        if ps is None or ps.get(q("val")) != '1':
            continue
        texts = []
        for t_elem in child.iter(q("t")):
            if t_elem.text:
                texts.append(t_elem.text)
        full = ''.join(texts).strip()
        if not full.startswith('第'):
            continue

        # Add page break before
        for pb in list(pPr.findall(q("pageBreakBefore"))):
            pPr.remove(pb)
        pb_elem = etree.SubElement(pPr, q("pageBreakBefore"))
        pb_elem.set(q("val"), "true")
        count += 1

    print(f"  为 {count} 个章标题添加了分页")

def main(args):
    SRC_DOCX = Path(args.input)
    OUT_DOCX = Path(args.output)

    print("=== 1. 解析序言.md ===")
    data = parse_md(MD_PATH)
    print(f"  序言: {len(data['preface_paragraphs'])} 段")
    for ch in data['chapters']:
        content_count = sum(1 for s in ch['sections'] if s['paragraphs'])
        print(f"  {ch['title']} ({len(ch['sections'])} 节, {content_count} 有正文)")

    print(f"\n=== 2. 加载 {SRC_DOCX.name} ===")
    doc = Document(str(SRC_DOCX))
    body = doc.element.body
    body_children = list(body)

    # ── Find key anchor elements ──

    # Chapter headings (Heading 1 style)
    chapter_heads = []  # (chapter_num, element, text)
    for child in body_children:
        if child.tag != q("p"):
            continue
        pPr = child.find(q("pPr"))
        if pPr is None:
            continue
        ps = pPr.find(q("pStyle"))
        if ps is None or ps.get(q("val")) != '1':
            continue
        texts = []
        for t in child.iter(q("t")):
            if t.text:
                texts.append(t.text)
        full = ''.join(texts).strip()
        m = re.match(r'第\s*(\d+)\s*章', full)
        if m:
            chapter_heads.append((int(m.group(1)), child, full))

    # 参考文献 and 致谢
    refs_elem = thanks_elem = None
    for child in reversed(list(body.iter(q("p")))):
        texts = []
        for t in child.iter(q("t")):
            if t.text:
                texts.append(t.text)
        full = ''.join(texts).strip()
        if thanks_elem is None and '致' in full and '谢' in full:
            thanks_elem = child
        if refs_elem is None and '参考文献' in full:
            refs_elem = child
        if refs_elem is not None and thanks_elem is not None:
            break

    # 序言 heading
    preface_head = None
    for child in body_children:
        if child.tag != q("p"):
            continue
        texts = []
        for t in child.iter(q("t")):
            if t.text:
                texts.append(t.text)
        full = ''.join(texts).strip()
        if '序' in full and '言' in full:
            preface_head = child
            break

    # Style template paragraphs
    sample_h2 = sample_h3 = None
    for child in body_children:
        if child.tag == q("p"):
            pPr = child.find(q("pPr"))
            if pPr is not None:
                ps = pPr.find(q("pStyle"))
                if ps is not None:
                    v = ps.get(q("val"))
                    if v == '2' and sample_h2 is None:
                        sample_h2 = child
                    elif v == '3' and sample_h3 is None:
                        sample_h3 = child
        if sample_h2 is not None and sample_h3 is not None:
            break

    print(f"  找到 {len(chapter_heads)} 个章标题")
    print(f"  序言标题: {'是' if preface_head is not None else '否'}")
    print(f"  参考文献: {'是' if refs_elem is not None else '否'}")

    # ── 3. Replace chapter headings & section structure ──
    print("\n=== 3. 替换章/节标题 ===")

    # Build index-based ranges for each chapter
    indexed = [(body_children.index(h[1]),) + h for h in chapter_heads]
    indexed.sort(key=lambda x: x[0])

    chapter_ranges = []
    for i, (idx, ch_num, ch_elem, ch_text) in enumerate(indexed):
        end_idx = indexed[i+1][0] if i+1 < len(indexed) else len(body_children)
        chapter_ranges.append({
            'num': ch_num,
            'start': idx,
            'end': end_idx,
            'element': ch_elem,
        })

    # Work backwards to avoid index drift
    for cr in reversed(chapter_ranges):
        ch_num = cr['num']
        ch_elem = cr['element']
        start = cr['start']
        end = cr['end']

        # Find matching chapter in md data
        md_ch = next((c for c in data['chapters'] if c['num'] == ch_num), None)
        if md_ch is None:
            print(f"  章{ch_num}: md中无匹配, 跳过")
            continue

        # Update chapter heading text
        texts = []
        for t in ch_elem.iter(q("t")):
            if t.text is not None:
                if not texts:
                    t.text = md_ch['title']
                else:
                    t.text = ''
                texts.append(t)
        _strip_sect_pr(ch_elem)

        # Remove all siblings between this chapter and next chapter
        siblings_to_remove = []
        sibling = ch_elem.getnext()
        while sibling is not None:
            if sibling.tag == q("p"):
                pPr = sibling.find(q("pPr"))
                if pPr is not None:
                    ps = pPr.find(q("pStyle"))
                    if ps is not None and ps.get(q("val")) == '1':
                        break
                if sibling in (refs_elem, thanks_elem):
                    break
            siblings_to_remove.append(sibling)
            sibling = sibling.getnext()

        for sib in siblings_to_remove:
            body.remove(sib)

        print(f"  章{ch_num}: 删除 {len(siblings_to_remove)} 个旧元素, 重建 {len(md_ch['sections'])} 节")

        # Insert new section headings + body text (in reverse for correct order)
        # Collect all new elements to insert
        insert_after = []
        for sec in reversed(md_ch['sections']):
            sec_text = f"{sec['number']} {sec['title']}"
            level = sec['level']
            style_val = str(level + 1)  # 1.x.y -> style '2', 1.x.y.z -> style '3'

            # Template selection
            template = None
            if level == 1 and sample_h2 is not None:
                template = sample_h2
            elif level == 2 and sample_h3 is not None:
                template = sample_h3

            heading_elem = _make_heading_para(sec_text, style_val, template)

            # If there's body content, process paragraphs into elements (detecting tables)
            body_elems = _process_paragraphs_to_elements(sec['paragraphs'])

            # Build insertion list (closest to heading first in reverse order)
            # Insert order (reverse): body paragraphs, then heading
            insert_after.append((heading_elem, body_elems, sec_text))

        for heading_elem, body_elems, sec_text in insert_after:
            # Insert body elements first (closest to heading), then heading
            # body_elems is now a list of (is_table, element) tuples
            for is_table, elem in reversed(body_elems):
                ch_elem.addnext(elem)
            ch_elem.addnext(heading_elem)
            if body_elems:
                para_count = sum(1 for is_t, _ in body_elems if not is_t)
                tbl_count = sum(1 for is_t, _ in body_elems if is_t)
                extra = f" (含 {tbl_count} 张表)" if tbl_count else ""
                print(f"    [{style_val}] {sec_text} + {para_count} 段正文{extra}")

    # ── 4. Handle preface content ──
    if preface_head is not None and data['preface_paragraphs']:
        print("\n=== 4. 插入序言正文 ===")
        # Remove old content after the 序言 heading up to the next heading-style paragraph
        siblings_to_remove = []
        sibling = preface_head.getnext()
        while sibling is not None:
            if sibling.tag == q("p"):
                pPr = sibling.find(q("pPr"))
                if pPr is not None:
                    ps = pPr.find(q("pStyle"))
                    if ps is not None and ps.get(q("val")) in ('1', '2'):
                        break
            siblings_to_remove.append(sibling)
            sibling = sibling.getnext()
        for sib in siblings_to_remove:
            body.remove(sib)

        body_elems = _process_paragraphs_to_elements(data['preface_paragraphs'])
        for is_table, elem in reversed(body_elems):
            preface_head.addnext(elem)
        print(f"  序言: 插入 {len(body_elems)} 个元素")

    # ── 5. Insert images from imgs/ into placeholder positions ──
    _insert_images(doc)

    # ── 6. Format references section per thesis standard ──
    _format_references(doc)

    # ── 7. Add page breaks before each chapter heading ──
    _add_chapter_page_breaks(doc)

    # ── 8. Save ──
    print(f"\n=== 5. 保存至 {OUT_DOCX.name} ===")
    try:
        doc.save(str(OUT_DOCX))
        print("完成！")
    except PermissionError:
        print(f"文件被占用, 保存至 {FALLBACK.name}")
        doc.save(str(FALLBACK))


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="将序言.md正文插入docx小标题下")
    parser.add_argument("--input", default=str(DEFAULT_SRC), help="输入docx路径")
    parser.add_argument("--output", default=str(DEFAULT_OUT), help="输出docx路径")
    args = parser.parse_args()
    main(args)
